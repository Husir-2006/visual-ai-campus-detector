from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE = Path(__file__).resolve().parent
DELIVERABLES = BASE / "deliverables"
CHARTS = BASE / "outputs" / "charts"
SAMPLES = BASE / "outputs" / "samples"


TRAIN_HISTORY = [
    (1, 0.2412, 0.1965),
    (2, 0.1375, 0.1668),
    (3, 0.1211, 0.1568),
    (4, 0.1030, 0.1533),
    (5, 0.0904, 0.1578),
    (6, 0.0773, 0.1482),
    (7, 0.0666, 0.1420),
    (8, 0.0583, 0.1486),
    (9, 0.0520, 0.1506),
    (10, 0.0475, 0.1503),
    (11, 0.0443, 0.1478),
    (12, 0.0425, 0.1511),
]


def font(size=24):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def make_training_chart():
    CHARTS.mkdir(parents=True, exist_ok=True)
    path = CHARTS / "training_curve.png"
    w, h = 1200, 640
    margin = 90
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_body = font(34), font(22)
    draw.text((margin, 28), "Training and validation loss", fill=(20, 30, 40), font=f_title)
    left, top, right, bottom = margin, 110, w - 70, h - 80
    draw.rectangle((left, top, right, bottom), outline=(210, 218, 228), width=2)
    for i in range(6):
        y = top + i * (bottom - top) / 5
        draw.line((left, y, right, y), fill=(235, 238, 243), width=1)
    max_y = 0.25

    def point(epoch, value):
        x = left + (epoch - 1) / 11 * (right - left)
        y = bottom - value / max_y * (bottom - top)
        return x, y

    train_points = [point(e, t) for e, t, _ in TRAIN_HISTORY]
    valid_points = [point(e, v) for e, _, v in TRAIN_HISTORY]
    draw.line(train_points, fill=(37, 99, 235), width=5)
    draw.line(valid_points, fill=(22, 134, 91), width=5)
    for p in train_points:
        draw.ellipse((p[0] - 5, p[1] - 5, p[0] + 5, p[1] + 5), fill=(37, 99, 235))
    for p in valid_points:
        draw.ellipse((p[0] - 5, p[1] - 5, p[0] + 5, p[1] + 5), fill=(22, 134, 91))
    draw.text((left, bottom + 24), "Epoch", fill=(80, 90, 105), font=f_body)
    draw.text((right - 260, bottom + 24), "train loss", fill=(37, 99, 235), font=f_body)
    draw.text((right - 130, bottom + 24), "valid loss", fill=(22, 134, 91), font=f_body)
    draw.text((left + 8, top + 8), "0.25", fill=(100, 110, 125), font=f_body)
    draw.text((left + 8, bottom - 30), "0.00", fill=(100, 110, 125), font=f_body)
    img.save(path)
    return path


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run(run, size=16 if level == 1 else 13, bold=True, color=(46, 116, 181))
    return paragraph


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(header)
        set_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(value))
            set_run(run)
    doc.add_paragraph()
    return table


def build_report():
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    chart = make_training_chart()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("基于 YOLO 的校园车辆与车牌检测系统设计与实现")
    set_run(r, size=20, bold=True, color=(11, 37, 69))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("创新应用综合实训课程报告")
    set_run(r, size=12, color=(80, 90, 105))

    add_heading(doc, "项目简要介绍", 1)
    add_para(doc, "本项目面向校园车辆管理与门禁辅助识别场景，设计并实现了一个车辆与车牌检测系统。系统支持图片上传、目标检测、检测框可视化、数量统计和车牌区域裁剪，形成了从模型训练到应用展示的完整流程。项目融合 PyTorch 轻量 YOLO 风格检测模型与 OpenCV 图像处理方法，并提供非静态网页系统和静态 HTML 演示页面。最终模型基于 25554 张图像训练，能够在测试集上完成车牌区域检测并输出可视化结果。")

    add_heading(doc, "成员工作描述及组内量化贡献比", 1)
    add_para(doc, "成员A负责项目选题、需求分析、报告撰写与整体材料整理；成员B负责数据集整理、模型训练脚本、训练参数调整与测试结果分析；成员C负责 Flask 后端、检测接口和 OpenCV 可视化模块；成员D负责前端页面、静态演示版、PPT 制作和演示视频准备。成员A：成员B：成员C：成员D = 2.5：3.0：2.5：2.0。")

    add_heading(doc, "一、项目背景与意义", 1)
    add_para(doc, "校园门口、停车场和道路区域经常需要对车辆通行情况进行记录与统计。传统人工查看监控画面的方式效率较低，也不利于后续数据分析。目标检测技术可以自动定位图片中的车辆和车牌区域，为车辆管理、通行统计和安防辅助提供技术基础。")

    add_heading(doc, "二、需求分析", 1)
    for item in [
        "支持上传车辆图片，并在页面中展示原图和检测结果图。",
        "能够检测车辆区域和车牌区域，并显示类别、置信度和坐标。",
        "能够统计车辆数量、车牌数量和平均置信度。",
        "能够裁剪车牌区域，便于课程报告和演示视频展示。",
        "提供非静态版和静态版两套入口，降低答辩展示风险。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、数据集与数据处理", 1)
    add_para(doc, "项目合并了三个公开车牌检测数据来源，均采用 YOLO 标注或可按 YOLO 目录结构读取。合并脚本会自动识别 train、valid、test 三个划分，并将图片与标签复制到统一目录。")
    add_table(
        doc,
        ["数据来源", "说明", "数量"],
        [
            ["Indonesian License Plate Detection Dataset", "Zenodo / Roboflow Universe，DOI: 10.5281/zenodo.15605718", "966"],
            ["Roboflow Public License Plates Dataset", "YOLO v5 PyTorch 格式", "350"],
            ["Kaggle ALPR Dataset", "Automatic License Plate Recognition 数据集", "24238"],
            ["合计", "合并后用于训练、验证、测试", "25554"],
        ],
    )
    add_table(
        doc,
        ["划分", "图片数量"],
        [["训练集", "22094"], ["验证集", "2310"], ["测试集", "1150"]],
    )

    add_heading(doc, "四、模型设计与训练", 1)
    add_para(doc, "考虑课程项目的可运行性和训练速度，本项目实现了一个轻量 YOLO 风格网格检测模型。模型将输入图像缩放到 256×256，并映射到 10×10 网格，每个网格预测目标置信度、边界框坐标和类别概率。训练时使用 AdamW 优化器、余弦退火学习率、Smooth L1 边界框损失和二分类交叉熵损失。")
    add_table(
        doc,
        ["训练配置", "取值"],
        [
            ["模型", "Tiny-YOLO 风格网格检测模型"],
            ["输入尺寸", "256×256"],
            ["网格尺寸", "10×10"],
            ["训练轮数", "12 epochs"],
            ["Batch size", "64"],
            ["数据增强", "亮度、对比度、噪声、轻微模糊"],
            ["最佳验证损失", "0.1420"],
            ["测试集损失", "0.1196"],
        ],
    )
    doc.add_picture(str(chart), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "五、系统设计与实现", 1)
    add_para(doc, "系统由前端展示层、后端接口层、模型检测层和结果输出层组成。非静态版使用 Flask 接收图片并调用训练好的 PyTorch 模型，同时结合 OpenCV 完成检测框绘制、车牌候选区域筛选和车牌裁剪。静态版使用单个 HTML 文件模拟完整演示流程，在无法启动后端时仍可完成答辩展示。")
    add_table(
        doc,
        ["模块", "主要功能"],
        [
            ["前端页面", "图片上传、原图预览、检测结果展示、统计指标展示"],
            ["检测接口", "接收图片、调用模型、返回 JSON 检测结果"],
            ["模型模块", "加载 tiny_plate_detector.pt 并输出候选框"],
            ["OpenCV 模块", "绘制检测框、裁剪车牌、补充车辆候选区域"],
            ["静态演示版", "无需服务器，直接打开 HTML 进行演示"],
        ],
    )

    add_heading(doc, "六、实验结果展示", 1)
    add_para(doc, "下图为测试集样例检测结果。系统能够输出车牌候选区域、置信度和坐标，并生成带标注的结果图。")
    sample = SAMPLES / "bigdata_sample_1.jpg"
    if sample.exists():
        doc.add_picture(str(sample), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "七、项目亮点与不足", 1)
    for item in [
        "使用 25554 张图像重新训练模型，相比初始小数据集更具有说服力。",
        "同时提供非静态版和静态版，兼顾真实模型推理和现场展示稳定性。",
        "系统保留 YOLOv5 ONNX 接口，后续可替换为更强的官方 YOLO 权重。",
        "当前模型偏向车牌区域检测，车辆整体检测主要由 OpenCV 候选区域补充，后续可加入完整车辆类别数据提升车辆检测能力。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "八、总结", 1)
    add_para(doc, "本项目完成了从数据集收集、数据合并、模型训练、后端服务、前端交互到演示材料整理的完整视觉 AI 应用流程。通过该项目，小组掌握了目标检测任务的基本数据格式、训练流程和系统落地方式，也为后续扩展实时视频检测、车牌字符识别和校园车辆管理系统奠定了基础。")

    out = DELIVERABLES / "课程报告-校园车辆与车牌检测系统.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build_report()
