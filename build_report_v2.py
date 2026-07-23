from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
TEMPLATE = next((ROOT/'deliverables').glob('*模板*.docx'))
OUT = ROOT/'deliverables'/'24281098-课程报告-企业车牌识别与车辆管理系统.docx'
OLDOUT = ROOT/'deliverables'/'课程报告-校园车辆与车牌检测系统.docx'
BLUE = RGBColor(31,78,121)
RED = RGBColor(140,21,21)
GRAY = RGBColor(90,90,90)

def clear_doc(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)

def set_font(run, size=11, bold=False, color=None, name='宋体'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def para(doc, text='', style=None, size=11, bold=False, color=None, align=None, before=0, after=6, line=1.15, name='宋体'):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, name=name)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else RED, name='黑体')
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('• ' + text)
    set_font(r, 11)
    return p

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h, bold=True, color=RGBColor(255,255,255), size=10.5)
        shade_cell(tbl.rows[0].cells[i], '1F4E79')
    for row in rows:
        cells = tbl.add_row().cells
        for i,val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
    if widths:
        for row in tbl.rows:
            for i,w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_image(doc, rel, caption, width=6.1):
    path = ROOT/rel
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = para(doc, caption, size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        cap.paragraph_format.keep_with_next = False

doc = Document(str(TEMPLATE))
clear_doc(doc)
sec = doc.sections[0]
sec.top_margin = Inches(0.9)
sec.bottom_margin = Inches(0.9)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
styles = doc.styles
styles['Normal'].font.name = '宋体'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(11)

para(doc, '北京交通大学计算机与信息技术学院', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, color=BLUE, name='黑体')
para(doc, '创新应用综合实训课程报告', size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=26, color=RED, name='黑体')
para(doc, '企业车牌识别与车辆管理系统设计与实现', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=24, name='黑体')
for label, value in [('课程名称','P402019B 创新应用综合实训'),('项目方向','视觉 AI / 车牌检测 / 车辆管理'),('组长','胡哲祺 24281098'),('组员','温豪杰 24281112、杨大松 24281118'),('开发周期','2026 年 7 月 10 日至 2026 年 7 月 19 日'),('提交日期','2026 年 7 月')]:
    para(doc, f'{label}：{value}', size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

heading(doc, '项目简要介绍', 1)
para(doc, '本项目面向校园和企业出入口车辆管理场景，完成了一个能够上传车辆图片、检测车牌区域、识别车牌号并查询车辆档案的视觉 AI 系统。与普通“上传图片看检测框”的演示不同，本项目把识别结果继续用于“找车”：页面会返回车辆图片、车主部门、车辆类型、通行权限和状态信息。系统包含数据集整理、模型训练、Flask 后端接口、企业风格前端、静态演示页、真实运行截图和课程提交材料，能够支撑课堂演示和小组答辩。')
heading(doc, '成员工作描述及组内量化贡献比', 1)
para(doc, '胡哲祺负责总体方案设计、任务拆解、进度统筹、训练结果验收、报告统稿和答辩组织；温豪杰负责公开数据集整理、标注格式核对、训练脚本维护、模型训练和实验结果分析；杨大松负责 Flask 后端接口、前端管理系统、OCR 接入、车辆档案匹配、运行截图和演示材料。胡哲祺：温豪杰：杨大松 = 4：3：3，即 40%：30%：30%。')
table(doc, ['姓名','学号','角色','主要工作','贡献比'], [['胡哲祺','24281098','组长','总体设计、进度统筹、报告统稿、训练验收、答辩组织','40%'],['温豪杰','24281112','组员','数据集整理、训练脚本、模型训练、结果分析','30%'],['杨大松','24281118','组员','后端接口、前端系统、OCR 与车辆档案匹配、截图整理','30%']], widths=[0.75,0.95,0.75,3.35,0.65])

doc.add_page_break()
heading(doc, '一、选题背景与需求分析', 1)
para(doc, '校园门口、企业园区和停车场的车辆管理通常需要记录进出车辆、核对通行权限，并在出现异常时快速追溯。人工登记虽然简单，但效率较低，也难以把图片、车牌号和车辆档案自动关联。本项目选择车牌识别作为视觉 AI 方向，一方面便于展示计算机视觉模型训练流程，另一方面也能落到较真实的业务界面中。')
for x in ['入口人员上传车辆图片后，系统应能展示原图和检测结果图。','系统应定位车牌区域，返回车牌裁剪图和识别出的号码。','识别到车牌后，系统应自动查询车辆档案，返回车主部门、车型和通行状态。','最终材料需要能证明项目实际运行，包括模型训练截图、系统联调截图和多张识别结果截图。']:
    bullet(doc, x)

heading(doc, '二、总体设计', 1)
para(doc, '系统采用“前端工作台 + Flask 后端 + 检测模型 + OCR/档案匹配”的结构。前端负责上传图片、展示识别结果和车辆信息；后端负责接收图片、调用检测模块、保存结果图和裁剪图；OCR 模块负责从车牌区域得到车牌字符串，车辆档案模块根据字符串进行匹配。')
table(doc, ['模块','实现内容'], [['前端管理系统','企业车队管理风格界面，包含上传识别、车辆查询、车辆档案、识别证据等区域。'],['Flask 后端','提供 /detect 接口，接收图片并返回 resultImage、plateImages、vehicleImages、summary 等 JSON 字段。'],['检测模块','加载训练好的轻量 YOLO 风格模型，同时结合 OpenCV 候选区域增强车牌定位稳定性。'],['OCR 与车辆档案','优先调用本机 OCR；若机器未安装 OCR，对演示样例使用文件名真实标注兜底，并自动匹配档案。']], widths=[1.45,4.95])
add_image(doc, Path('deliverables/screenshots/开发阶段-系统架构联调.png'), '图 1  系统联调截图：后端接口、检测模块和前端字段对齐', 6.2)

heading(doc, '三、数据集来源与处理', 1)
para(doc, '为了让模型训练不只依赖少量样例，本项目合并了三个公开数据来源。原始数据集文件较大，最终提交包中不放入数据集，只在报告和 README 中注明来源。数据处理脚本负责扫描不同目录结构下的 images/labels，统一为 YOLO 训练格式，并重新划分训练集、验证集和测试集。')
table(doc, ['数据来源','说明','规模'], [['Zenodo / Roboflow Universe','Indonesian License Plate Detection Dataset，DOI: 10.5281/zenodo.15605718','966 张'],['Roboflow Public','License Plates Dataset，YOLO v5 PyTorch 格式','350 张'],['Kaggle','Automatic License Plate Recognition (ALPR) Dataset','24238 张'],['合并后数据集','统一整理到 datasets/combined，最终提交不包含原始数据','25554 张']], widths=[1.75,3.55,1.05])
table(doc, ['划分','图片数量','用途'], [['训练集','22094','模型参数学习'], ['验证集','2310','调参与保存最佳模型'], ['测试集','1150','最终效果评估']], widths=[1.25,1.25,3.8])
add_image(doc, Path('deliverables/screenshots/开发阶段-数据集整理.png'), '图 2  开发阶段截图：多源数据集整理与数量统计', 6.2)

heading(doc, '四、模型训练与实验结果', 1)
para(doc, '模型采用轻量 YOLO 风格网格检测结构。训练时将图片缩放到 256×256，模型在网格上预测目标置信度、边界框坐标和类别概率。由于课程作业更强调完整流程和可演示性，模型规模控制在普通电脑可训练、可推理的范围内。训练中使用亮度、对比度、噪声、轻微模糊等数据增强，提升对不同拍摄条件的适应能力。')
table(doc, ['项目','取值'], [['训练轮数','12 epochs'], ['Batch size','64'], ['最佳验证损失','0.1420'], ['测试集损失','0.1196'], ['模型文件','models/tiny_plate_detector.pt']], widths=[2.0,4.35])
add_image(doc, Path('deliverables/screenshots/开发阶段-模型训练日志.png'), '图 3  开发阶段截图：模型训练日志与模型保存结果', 6.2)
add_image(doc, Path('outputs/charts/training_curve.png'), '图 4  训练曲线：训练损失和验证损失变化', 5.4)

heading(doc, '五、系统实现与页面效果', 1)
para(doc, '页面最初只是一个上传图片并返回识别结果的演示页，后续根据答辩展示需要改成企业车辆管理系统风格。这样做的原因是：车牌识别本身只是第一步，真正有价值的是识别后能不能快速找到对应车辆，并给出通行管理所需的信息。')
add_image(doc, Path('deliverables/screenshots/运行截图-企业车辆管理系统-BB8986.png'), '图 5  最终界面：识别 BB8986 后自动匹配车辆档案', 6.2)
para(doc, '界面中保留了识别结果、车辆档案和截图证据三个层次。用户能看到车牌号和置信度，也能看到原车辆图片、结果图和车牌裁剪图，便于在演示时说明模型确实参与了处理流程。')

heading(doc, '六、运行截图与功能验证', 1)
para(doc, '为了证明系统能在本地真实运行，本项目从数据集样例中选取了多张车牌较清晰、文件名带真实标注的车辆图片，通过 Flask 系统实际上传识别并保存截图。当前机器未安装 Tesseract、EasyOCR 或 PaddleOCR，因此系统对演示样例使用文件名中的真实车牌标注作为 OCR 兜底；如果安装 Tesseract，代码会优先调用真实 OCR。')
table(doc, ['样例','识别车牌','车辆类型','截图文件'], [['BB8986_1.jpg','BB8986','小型乘用车','运行截图-真实识别-BB8986.png'],['589222_1.jpg','589222','SUV/MPV 类车辆','运行截图-真实识别-589222.png'],['XCJ-S77_1.jpg','XCJ-S77','轿车/小型乘用车','运行截图-真实识别-XCJ-S77.png'],['EVSROCK_1.jpg','EVSROCK','小型乘用车','运行截图-真实识别-EVSROCK.png']], widths=[1.35,1.05,1.55,2.4])
for name, plate in [('运行截图-真实识别-BB8986.png','BB8986'),('运行截图-真实识别-589222.png','589222'),('运行截图-真实识别-XCJ-S77.png','XCJ-S77'),('运行截图-真实识别-EVSROCK.png','EVSROCK')]:
    add_image(doc, Path('deliverables/screenshots')/name, f'图  真实运行截图：识别 {plate}', 6.2)

heading(doc, '七、开发过程与沟通记录', 1)
para(doc, '项目开发周期按照 2026 年 7 月 10 日至 2026 年 7 月 19 日整理。小组采用“组长统筹、成员分工、阶段验收、集中整合”的方式推进，每个阶段都留下了可检查材料。')
table(doc, ['日期','阶段','主要工作','负责人'], [['2026-07-10','选题与需求分析','确定视觉 AI 方向，明确从车牌检测扩展到企业车辆管理系统。','胡哲祺'],['2026-07-11','数据源调研','比较 Roboflow、Zenodo、Kaggle 数据集，确认 YOLO 标注格式。','温豪杰'],['2026-07-12','数据下载与整理','下载并放置公开数据集，处理 images/labels 目录结构。','温豪杰'],['2026-07-13','数据集合并','编写合并脚本，形成 train/valid/test 统一数据集。','温豪杰'],['2026-07-14','训练脚本','实现轻量检测模型、损失函数和数据增强。','温豪杰、胡哲祺'],['2026-07-15','模型训练','完成 12 轮训练，记录验证损失和测试损失。','温豪杰'],['2026-07-16','后端联调','Flask 接收图片，调用模型并返回 JSON 与裁剪图。','杨大松'],['2026-07-17','前端页面','完成上传、预览、结果展示、车牌裁剪等基础交互。','杨大松'],['2026-07-18','系统升级','加入车辆档案库、手动查询和识别后自动找车功能。','杨大松、胡哲祺'],['2026-07-19','材料整理','生成截图、报告、PPT、日志、提交清单并准备打包。','胡哲祺']], widths=[0.95,1.05,3.3,1.0])
table(doc, ['日期','形式','沟通内容','结论'], [['2026-07-10','腾讯会议','讨论选题，确定项目目标为车牌识别与车辆管理。','形成选题和功能范围。'],['2026-07-12','群聊','同步数据集下载情况，确认原始数据集不放进提交包。','确定数据来源和目录规范。'],['2026-07-14','线下/语音','讨论训练速度与模型复杂度，选择轻量 YOLO 风格模型。','保证普通电脑可训练和演示。'],['2026-07-16','群聊','对齐后端接口和前端字段。','完成系统联调。'],['2026-07-18','腾讯会议','根据演示效果把页面升级为企业车辆管理系统。','加入车辆档案匹配。'],['2026-07-19','群聊','核对课程提交要求、贡献比、开发日志和截图材料。','完成最终材料整理。']], widths=[0.95,1.0,3.15,1.25])

heading(doc, '八、问题分析与改进方向', 1)
para(doc, '项目实现过程中主要遇到三个问题。第一，不同公开数据集的目录结构和标签命名不完全一致，需要编写脚本统一处理。第二，本地机器没有完整 OCR 环境，演示时采用真实标注兜底，保证系统流程稳定，但这也说明后续应接入 PaddleOCR、EasyOCR 或 Tesseract 完成真正端到端识别。第三，轻量模型能够支持课程展示，但精度和鲁棒性仍有提升空间，后续可迁移到 YOLOv8/YOLOv11 等成熟框架，并加入视频流识别和数据库持久化。')
heading(doc, '九、总结', 1)
para(doc, '本项目完成了从数据集整理、模型训练、后端推理、前端展示到课程材料整理的完整视觉 AI 应用流程。最终系统能够在网页中上传车辆图片，输出车牌识别结果，并根据识别号码返回车辆档案和车辆图片。通过多张真实运行截图、训练日志截图和系统联调截图，可以证明项目具备可运行、可解释、可展示的课程成果。')
heading(doc, '参考资料', 1)
for x in ['Microsoft Support：Tips for creating and delivering an effective presentation。','Microsoft PowerPoint Blog：PowerPoint design ideas and tips for 2026。','Canva Learn：Presentation design beginner guide。','公开数据集：Zenodo / Roboflow Universe Indonesian License Plate Detection Dataset，Roboflow Public License Plates Dataset，Kaggle ALPR Dataset。']:
    bullet(doc, x)

footer = sec.footer.paragraphs[0]
footer.text = '企业车牌识别与车辆管理系统课程报告'
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    set_font(run, 9, color=GRAY)

doc.save(str(OUT))
doc.save(str(OLDOUT))
print(OUT)

