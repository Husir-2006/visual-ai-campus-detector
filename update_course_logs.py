from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"G:\Grade2SummerSemester\InnovationApplication\visual-ai-campus-detector")
DELIVERABLES = PROJECT / "deliverables"

TEAM = [
    ("胡哲祺", "24281098", "组长", "总体设计、进度统筹、报告统稿、模型训练验收、答辩组织", "40%"),
    ("温豪杰", "24281112", "组员", "数据集整理、训练脚本、模型训练、实验结果分析", "30%"),
    ("杨大松", "24281118", "组员", "Flask 后端、管理系统前端、OCR 与车辆档案匹配、运行截图和演示材料", "30%"),
]

SCHEDULE = [
    ("2026-07-10", "选题与需求分析", "确定视觉 AI 方向，将目标从车牌检测扩展到车辆档案查询。", "胡哲祺"),
    ("2026-07-11", "数据源调研", "比较 Zenodo、Roboflow、Kaggle 数据集及标注格式。", "温豪杰"),
    ("2026-07-12", "数据下载与整理", "完成公开数据下载和 images/labels 目录整理。", "温豪杰"),
    ("2026-07-13", "数据集合并", "统一 YOLO 标注格式，生成 train/valid/test 数据集。", "温豪杰"),
    ("2026-07-14", "训练脚本实现", "完成轻量检测模型、损失函数和数据增强。", "温豪杰、胡哲祺"),
    ("2026-07-15", "模型训练与调参", "完成 12 轮训练，保存最佳权重并记录实验指标。", "温豪杰"),
    ("2026-07-16", "后端接口联调", "完成 Flask 图片上传、模型调用和 JSON 结果返回。", "杨大松"),
    ("2026-07-17", "前端识别页面", "完成上传、预览、检测结果和车牌裁剪图展示。", "杨大松"),
    ("2026-07-18", "管理系统升级", "加入车辆档案、按牌找车、权限与状态信息。", "杨大松、胡哲祺"),
    ("2026-07-19", "材料整理与验收", "生成运行截图，更新报告、PPT、日志、讲稿和提交包。", "胡哲祺"),
]

COMMUNICATIONS = [
    ("2026-07-10", "腾讯会议", "讨论课程要求和视觉 AI 选题。", "确定企业车牌识别与车辆管理方向。"),
    ("2026-07-12", "群聊", "同步数据下载进度和目录结构。", "原始数据集不进入提交包，只注明来源。"),
    ("2026-07-13", "群聊", "核对三套数据的标注格式和合并脚本。", "统一为 YOLO 训练目录。"),
    ("2026-07-14", "线下/语音", "讨论模型复杂度、训练速度和演示需求。", "采用普通电脑可训练的轻量模型。"),
    ("2026-07-16", "群聊", "对齐后端接口和前端字段。", "统一结果图、车牌图、车辆图和 summary 字段。"),
    ("2026-07-17", "线下讨论", "检查上传、预览、结果图和裁剪图流程。", "确认前端第一版可用。"),
    ("2026-07-18", "腾讯会议", "复盘演示效果，讨论如何体现业务价值。", "升级为企业车辆管理工作台并加入自动找车。"),
    ("2026-07-19", "群聊", "核对分工、日期、截图和命名要求。", "按 40%/30%/30% 完成最终材料。"),
]

PERSONAL = {
    "胡哲祺": [
        ("2026-07-10", "组织选题讨论，梳理课程要求，确定企业车牌识别与车辆管理方向。"),
        ("2026-07-11", "审核公开数据集候选方案，确定三个数据源及最终成果范围。"),
        ("2026-07-12", "检查数据下载进度和目录规范，明确原始数据集不进入提交包。"),
        ("2026-07-13", "验收数据合并结果，确认训练、验证和测试集数量。"),
        ("2026-07-14", "参与模型结构和训练参数讨论，确定轻量训练方案。"),
        ("2026-07-15", "检查训练日志、损失变化和权重文件，记录实验关键指标。"),
        ("2026-07-16", "组织前后端接口联调，核对返回字段和截图证据需求。"),
        ("2026-07-17", "验收前端第一版，提出增加车辆档案和自动找车功能。"),
        ("2026-07-18", "统筹企业管理界面升级，核对演示逻辑和课程亮点。"),
        ("2026-07-19", "统稿报告、PPT、日志和讲稿，完成材料检查与打包。"),
    ],
    "温豪杰": [
        ("2026-07-10", "参与选题讨论，提出以公开车牌检测数据开展视觉 AI 实验。"),
        ("2026-07-11", "调研 Zenodo、Roboflow 和 Kaggle 数据集，比较规模与标注格式。"),
        ("2026-07-12", "下载并检查数据文件，整理 images、labels、train、valid、test 目录。"),
        ("2026-07-13", "编写和运行数据集合并脚本，处理重复文件名与标签路径。"),
        ("2026-07-14", "实现轻量检测模型训练脚本、数据增强和损失计算。"),
        ("2026-07-15", "完成 12 轮训练，记录验证损失、测试损失并保存最佳权重。"),
        ("2026-07-16", "说明模型输入输出格式，协助后端验证推理结果。"),
        ("2026-07-17", "抽取测试样例复核检测框和车牌裁剪结果，整理问题清单。"),
        ("2026-07-18", "补充模型与数据说明，核对 PPT 和报告中的实验数据。"),
        ("2026-07-19", "复查数据来源、训练参数和实验结论，协助最终验收。"),
    ],
    "杨大松": [
        ("2026-07-10", "参与需求讨论，提出识别后返回车辆图片和档案信息的展示方式。"),
        ("2026-07-11", "调研 Flask 图片上传和前端结果展示方案，整理接口草案。"),
        ("2026-07-12", "搭建基础项目目录和静态页面，准备图片上传区域。"),
        ("2026-07-13", "实现结果图、车牌裁剪图和车辆图片的文件访问逻辑。"),
        ("2026-07-14", "配合训练模块确定推理调用方式，准备后端模型加载接口。"),
        ("2026-07-15", "接入模型权重，完成单张图片推理和 JSON 返回测试。"),
        ("2026-07-16", "完成 Flask /detect 接口和前后端字段联调。"),
        ("2026-07-17", "完成上传、预览、识别结果和裁剪图展示的前端第一版。"),
        ("2026-07-18", "升级为企业车辆管理界面，加入车辆档案和自动找车功能。"),
        ("2026-07-19", "运行多组样例并整理截图，完善静态演示版和演示材料。"),
    ],
}


def set_font(run, size=10.5, bold=False, color="202124"):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, value, bold=False, color="202124", size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value))
    set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        shade(table.rows[0].cells[index], "F1F3F4")
        set_cell(table.rows[0].cells[index], header, bold=True, size=9.5)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell(cells[index], value, size=9.2)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
    return table


def build_personal_log(name, sid, role, focus):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("北京交通大学计算机与信息技术学院\n实习实训日志")
    set_font(run, size=16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run(f"{name}（{sid}）  |  {role}  |  2026-07-10 至 2026-07-19")
    set_font(run, size=10.5, color="5F6368")

    rows = [(date, content, "已完成") for date, content in PERSONAL[name]]
    add_table(doc, ["日期", "工作内容", "完成情况"], rows, [1.05, 4.65, 0.9])

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(12)
    summary.paragraph_format.line_spacing = 1.25
    run = summary.add_run(
        f"个人总结：本阶段主要承担{focus}。通过连续 10 天的开发和协作，"
        "进一步熟悉了视觉 AI 项目从数据、模型到后端接口、前端展示和课程材料整理的完整过程。"
    )
    set_font(run, size=10.5)

    output = DELIVERABLES / f"北京交通大学计算机与信息技术学院实习实训日志-{name}-{sid}.docx"
    doc.save(output)
    return output


def build_markdown():
    lines = [
        "# 开发日志与沟通交流记录",
        "",
        "开发周期：2026-07-10 至 2026-07-19",
        "",
        "说明：以下内容根据代码、训练记录、运行截图和小组讨论过程整理。",
        "",
        "## 成员与贡献",
        "",
    ]
    for name, sid, role, focus, contribution in TEAM:
        lines.append(f"- {role} {name}（{sid}）：{focus}，贡献度 {contribution}。")
    lines.extend([
        "",
        "贡献比：胡哲祺：温豪杰：杨大松 = 4：3：3，即 40%：30%：30%。",
        "",
        "## 开发日志",
        "",
        "| 日期 | 阶段 | 工作内容 | 负责人 |",
        "|---|---|---|---|",
    ])
    for row in SCHEDULE:
        lines.append("| " + " | ".join(row) + " |")
    lines.extend([
        "",
        "## 沟通交流记录",
        "",
        "| 日期 | 形式 | 沟通内容 | 结论 |",
        "|---|---|---|---|",
    ])
    for row in COMMUNICATIONS:
        lines.append("| " + " | ".join(row) + " |")
    content = "\n".join(lines) + "\n"
    for filename in ["开发日志与沟通交流记录.md", "24281098-开发日志与沟通交流记录.md"]:
        (DELIVERABLES / filename).write_text(content, encoding="utf-8")


def main():
    for name, sid, role, focus, _ in TEAM:
        print(build_personal_log(name, sid, role, focus))
    build_markdown()


if __name__ == "__main__":
    main()
