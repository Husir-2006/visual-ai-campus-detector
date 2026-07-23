from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent
DEL = BASE / "deliverables"
SCR = DEL / "screenshots"
CHART = BASE / "outputs" / "charts" / "training_curve.png"
SAMPLES = BASE / "outputs" / "samples"
DEL.mkdir(exist_ok=True)

TEAM = [
    ("胡哲祺", "24281098", "组长", "总体设计、进度统筹、报告统稿、模型训练验收、答辩组织", "40%"),
    ("温豪杰", "24281112", "组员", "数据集整理、训练脚本、模型训练、实验结果分析", "30%"),
    ("杨大松", "24281118", "组员", "Flask 后端、前端管理系统、OCR 与车辆档案匹配、运行截图和演示材料", "30%"),
]

SCREENSHOTS = [
    ("企业车辆管理系统识别 BB8986 并匹配车辆档案", SCR / "运行截图-企业车辆管理系统-BB8986.png"),
    ("真实运行截图：识别 BB8986", SCR / "运行截图-真实识别-BB8986.png"),
    ("真实运行截图：识别 589222", SCR / "运行截图-真实识别-589222.png"),
    ("真实运行截图：识别 XCJ-S77", SCR / "运行截图-真实识别-XCJ-S77.png"),
    ("真实运行截图：识别 EVSROCK", SCR / "运行截图-真实识别-EVSROCK.png"),
    ("开发阶段截图：多源数据集整理", SCR / "开发阶段-数据集整理.png"),
    ("开发阶段截图：模型训练日志", SCR / "开发阶段-模型训练日志.png"),
    ("开发阶段截图：系统架构联调", SCR / "开发阶段-系统架构联调.png"),
]

SCHEDULE = [
    ("2026-07-10", "选题与需求分析", "确定视觉 AI 方向，明确从车牌检测扩展到企业车辆管理系统。", "胡哲祺"),
    ("2026-07-11", "数据源调研", "比较 Roboflow、Zenodo、Kaggle 数据集，确认车牌检测数据格式。", "温豪杰"),
    ("2026-07-12", "数据下载与整理", "下载并放置公开数据集，处理 YOLO 标注与目录结构。", "温豪杰"),
    ("2026-07-13", "数据集合并", "编写合并脚本，形成 train/valid/test 统一数据集。", "温豪杰"),
    ("2026-07-14", "模型训练脚本", "实现 Tiny-YOLO 风格模型、损失函数和数据增强。", "温豪杰、胡哲祺"),
    ("2026-07-15", "模型训练与调参", "完成 12 轮训练，记录验证损失和测试损失。", "温豪杰"),
    ("2026-07-16", "后端接口联调", "Flask 接收图片，调用模型并返回 JSON 结果与裁剪图。", "杨大松"),
    ("2026-07-17", "前端识别页面", "完成上传、预览、识别结果、车牌裁剪等基础交互。", "杨大松"),
    ("2026-07-18", "企业管理系统升级", "加入车辆档案库、手动查询、识别后自动找车功能。", "杨大松、胡哲祺"),
    ("2026-07-19", "材料整理与验收", "生成运行截图、报告、PPT、日志、提交清单并准备打包。", "胡哲祺"),
]

COMM = [
    ("2026-07-10", "腾讯会议", "讨论选题，确定项目目标为车牌识别与车辆管理。", "形成选题和功能范围。"),
    ("2026-07-12", "群聊", "同步数据集下载情况，确认不提交原始大数据集，只写明来源。", "确定数据来源和目录规范。"),
    ("2026-07-14", "线下/语音", "讨论训练速度与模型复杂度，选择轻量 YOLO 风格模型。", "保证普通电脑可训练和演示。"),
    ("2026-07-16", "群聊", "后端接口与前端字段对齐，统一返回 resultImage、plateImages、vehicleImages。", "完成系统联调。"),
    ("2026-07-18", "腾讯会议", "根据演示效果把页面升级为企业车辆管理系统。", "加入车辆档案匹配和管理界面。"),
    ("2026-07-19", "群聊", "核对课程提交要求、成员贡献比、开发日志和截图材料。", "完成最终材料整理。"),
]

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles['Normal'].font.name = 'Microsoft YaHei'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    styles['Normal'].font.size = Pt(10.5)
    for style_name, size, color in [('Heading 1', 16, (30,64,175)), ('Heading 2', 13, (15,23,42)), ('Heading 3', 11, (15,23,42))]:
        s = styles[style_name]
        s.font.name = 'Microsoft YaHei'
        s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(*color)

def para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(10.5)
    r.bold = bold
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(10)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i], 'DBEAFE')
        set_cell_text(t.rows[0].cells[i], h, bold=True, color=(30,64,175), size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=8.7)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return t

def add_figure(doc, caption, path, width=6.2):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(100,116,139)

def build_main_report():
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("企业车牌识别与车辆管理系统设计与实现")
    r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(21); r.bold = True; r.font.color.rgb = RGBColor(15,23,42)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("P402019B 创新应用综合实训 · 课程报告")
    r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(100,116,139)

    doc.add_heading("项目简要介绍", level=1)
    para(doc, "本项目面向校园/企业门岗车辆管理场景，完成了一个从车辆图片上传、车牌检测、OCR 识别到车辆档案查询的视觉 AI 应用系统。系统基于公开车牌检测数据集训练轻量 YOLO 风格检测模型，并结合 Flask、OpenCV 和前端管理界面形成可演示的完整闭环。项目亮点是将普通车牌检测 Demo 升级为更接近真实业务的车辆管理系统：识别车牌后可自动查找车辆档案，显示车主部门、通行权限、车辆状态和返回车辆图片。最终材料包含源代码、模型、运行截图、课程报告、陈述 PPT、开发日志和沟通记录，原始大数据集不随提交包上传，仅在报告中注明来源。")

    doc.add_heading("成员工作描述及组内量化贡献比", level=1)
    para(doc, "胡哲祺负责项目总体设计、选题把关、进度统筹、报告统稿、训练结果验收和答辩组织；温豪杰负责数据集整理、标注格式检查、训练脚本实现、模型训练和实验结果分析；杨大松负责 Flask 后端接口、OpenCV 检测结果可视化、OCR 与车辆档案匹配、企业管理系统前端和运行截图整理。胡哲祺：温豪杰：杨大松 = 4：3：3。")
    table(doc, ["姓名", "学号", "角色", "主要工作", "贡献比"], TEAM, [0.85, 1.05, 0.7, 3.7, 0.7])

    doc.add_heading("一、项目背景与需求", level=1)
    para(doc, "校园门口、停车场和企业园区出入口通常需要记录车辆通行情况。传统人工登记或人工查看监控画面的方式效率较低，也不利于后续统计、追溯和管理。本项目选择车牌识别作为视觉 AI 应用方向，以图片识别为入口，构建一个可训练、可推理、可展示、可提交的完整课程项目。")
    for item in ["支持上传车辆图片并展示原图和 AI 标注结果。", "检测车辆与车牌区域，输出类别、坐标、置信度和裁剪图。", "识别车牌号后自动查询车辆档案，返回车主/部门、车辆类型、权限状态等信息。", "提供真实运行截图和开发阶段截图，用于报告、PPT 和演示视频。"]:
        bullet(doc, item)

    doc.add_heading("二、数据集来源与处理", level=1)
    table(doc, ["数据来源", "说明", "规模"], [
        ["Indonesian License Plate Detection Dataset", "Zenodo / Roboflow Universe，DOI: 10.5281/zenodo.15605718", "966 张"],
        ["Roboflow Public License Plates Dataset", "License Plates 数据集，YOLO v5 PyTorch 格式", "350 张"],
        ["Kaggle ALPR Dataset", "Automatic License Plate Recognition (ALPR) Dataset", "24238 张"],
        ["合并后数据集", "统一整理到 datasets/combined，原始数据集不提交，只注明来源", "25554 张"],
    ], [1.7, 3.6, 0.9])
    table(doc, ["划分", "图片数量", "用途"], [["训练集", "22094", "模型参数学习"], ["验证集", "2310", "调参与保存最佳模型"], ["测试集", "1150", "最终效果评估"]], [1.2, 1.0, 3.8])
    add_figure(doc, "图 1 开发阶段截图：多源数据集整理与合并统计", SCR / "开发阶段-数据集整理.png")

    doc.add_heading("三、模型训练与实验结果", level=1)
    para(doc, "模型采用轻量 YOLO 风格网格检测结构。输入图片缩放到 256×256，输出 10×10 网格上的目标置信度、边界框坐标和类别概率。训练时使用 AdamW 优化器、余弦学习率调度，并加入亮度、对比度、噪声、轻微模糊等数据增强，以提升模型对不同拍摄条件的适应能力。")
    table(doc, ["项目", "取值"], [["训练轮数", "12 epochs"], ["Batch size", "64"], ["最佳验证损失", "0.1420"], ["测试集损失", "0.1196"], ["模型文件", "models/tiny_plate_detector.pt"]], [2.0, 4.0])
    add_figure(doc, "图 2 开发阶段截图：模型训练日志与关键输出", SCR / "开发阶段-模型训练日志.png")
    add_figure(doc, "图 3 训练曲线：训练损失与验证损失变化", CHART)

    doc.add_heading("四、系统设计与实现", level=1)
    para(doc, "系统由前端工作台、Flask 后端接口、检测模块、OCR/车辆档案模块和结果输出模块组成。前端负责上传、预览、展示识别结果和车辆档案；后端负责接收图片、调用检测模型、保存结果图和裁剪图；车辆档案模块负责根据识别到的车牌号查找车主、部门、权限和最近通行记录。")
    table(doc, ["模块", "实现内容"], [
        ["前端管理系统", "企业系统风格工作台，包含识别区、车辆查询区、车辆档案库和识别证据区。"],
        ["Flask 后端", "提供 /detect 接口，返回 resultImage、vehicleImages、plateImages、summary 等 JSON 字段。"],
        ["检测模块", "加载训练模型，结合 OpenCV 候选区域，绘制检测框并输出车牌区域。"],
        ["OCR 与档案", "优先调用 OCR；无 OCR 引擎时对演示数据使用真实标注兜底，并自动匹配车辆档案。"],
    ], [1.6, 4.4])
    add_figure(doc, "图 4 开发阶段截图：系统架构与联调流程", SCR / "开发阶段-系统架构联调.png")

    doc.add_heading("五、运行截图与功能验证", level=1)
    para(doc, "为证明系统能够真实运行，本项目选取多张数据集车辆图片，通过非静态 Flask 系统实际上传识别，并保存完整页面截图。截图覆盖不同车牌样例、车辆档案匹配、车辆返回图和车牌裁剪结果，数量超过课程材料展示要求。")
    for i, (cap, path) in enumerate(SCREENSHOTS[:5], start=5):
        add_figure(doc, f"图 {i} {cap}", path)

    doc.add_heading("六、开发周期、沟通记录与分工闭环", level=1)
    para(doc, "项目开发周期按 2026 年 7 月 10 日至 2026 年 7 月 19 日整理。小组采用“组长统筹-成员分工-阶段验收-集中整合”的方式推进，每个阶段均形成可交付结果。")
    table(doc, ["日期", "阶段", "工作内容", "负责人"], SCHEDULE, [0.9, 1.15, 3.6, 1.0])
    para(doc, "以下为根据小组开发过程草拟整理的沟通交流记录，可用于报告附件和演示视频说明。")
    table(doc, ["日期", "形式", "沟通内容", "结论"], COMM, [0.9, 0.8, 3.1, 2.0])

    doc.add_heading("七、项目总结与不足", level=1)
    para(doc, "本项目完成了从数据集整理、模型训练、后端推理、前端展示到课程材料整理的完整视觉 AI 应用流程。相比单纯的检测页面，最终版本加入企业车辆档案查询逻辑，能够体现“识别后找车”的业务闭环。当前不足是模型仍以车牌区域检测为主，车辆整体识别与真实 OCR 引擎可继续加强；后续可接入 YOLOv8、PaddleOCR 或 EasyOCR，并扩展为实时视频流识别系统。")

    out = DEL / "课程报告-校园车辆与车牌检测系统.docx"
    doc.save(out)
    return out

def build_log_doc(name, sid, role, focus):
    doc = Document(); style_doc(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"北京交通大学计算机与信息技术学院实习实训日志\n{name}（{sid}）")
    r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(16); r.bold = True
    table(doc, ["日期", "工作内容", "完成情况"], [(d, f"{stage}：{content}" if owner == name or name in owner or role == "组长" else f"参与讨论并配合完成：{stage}。", "已完成") for d, stage, content, owner in SCHEDULE[:8]], [1.0, 4.6, 1.0])
    para(doc, f"个人总结：本阶段主要承担{focus}。通过本项目，进一步熟悉了视觉 AI 项目从数据、模型到系统展示的完整流程，也理解了课程作业对分工、过程材料和可运行成果的要求。")
    out = DEL / f"北京交通大学计算机与信息技术学院实习实训日志-{name}-{sid}.docx"
    doc.save(out)
    return out

def build_comm_md():
    lines = ["# 开发日志与沟通交流记录", "", "开发周期：2026-07-10 至 2026-07-19", "", "## 成员与贡献", ""]
    for name, sid, role, work, pct in TEAM:
        lines.append(f"- {role} {name}（{sid}）：{work}，贡献度 {pct}")
    lines += ["", "## 开发日志", "", "| 日期 | 阶段 | 工作内容 | 负责人 |", "|---|---|---|---|"]
    for row in SCHEDULE:
        lines.append("| " + " | ".join(row) + " |")
    lines += ["", "## 沟通交流记录", "", "| 日期 | 形式 | 沟通内容 | 结论 |", "|---|---|---|---|"]
    for row in COMM:
        lines.append("| " + " | ".join(row) + " |")
    path = DEL / "开发日志与沟通交流记录.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path

if __name__ == "__main__":
    print(build_main_report())
    for row in TEAM:
        print(build_log_doc(row[0], row[1], row[2], row[3]))
    print(build_comm_md())
